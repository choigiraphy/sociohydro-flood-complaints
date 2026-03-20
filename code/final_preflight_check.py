from __future__ import annotations

from pathlib import Path
import sys

import pandas as pd
from docx import Document


DEFAULT_DOCX_PATH = Path(
    "/Volumes/Keywest_JetDrive 1/학교/2026-1/AGU_WRR_Choi/최신_AGU_Manuscript_Choi_수정중_Level2_20260320_v10.docx"
)
REQUIRED_FILES = [
    Path("/Volumes/Keywest_JetDrive 1/Codex/Academia/Complaint/tables/bootstrap_results.xlsx"),
    Path("/Volumes/Keywest_JetDrive 1/Codex/Academia/Complaint/tables/baseline_distribution_summary.xlsx"),
    Path("/Volumes/Keywest_JetDrive 1/Codex/Academia/Complaint/tables/threshold_sensitivity.xlsx"),
    Path("/Volumes/Keywest_JetDrive 1/Codex/Academia/Complaint/figures/Figure5_WRR.png"),
    Path("/Volumes/Keywest_JetDrive 1/Codex/Academia/Complaint/figures/Figure6_WRR.png"),
]

FORBIDDEN_STRINGS = [
    "±30 days",
    "rolling 14-day",
    "2-week rolling complaint sum",
    "Inistitutionalized",
    "⁴",
    "⁵",
]

EXPECTED_SUPP_REFS = [
    "Supplementary Table S1",
    "Supplementary Figure S1",
    "Supplementary Figure S2",
    "Supplementary Figure S3",
    "Supplementary Table S2",
    "Supplementary Figure S4",
    "Supplementary Table S5",
]


def ok(msg: str) -> None:
    print(f"[OK] {msg}")


def fail(msg: str) -> None:
    print(f"[FAIL] {msg}")


def warn(msg: str) -> None:
    print(f"[WARN] {msg}")


def check_files(docx_path: Path) -> list[str]:
    errors: list[str] = []
    if not docx_path.exists():
        errors.append(f"Missing required file: {docx_path}")
    else:
        ok(f"Found {docx_path.name}")
    for path in REQUIRED_FILES:
        if not path.exists():
            errors.append(f"Missing required file: {path}")
        else:
            ok(f"Found {path.name}")
    return errors


def load_doc(docx_path: Path) -> Document:
    return Document(str(docx_path))


def check_forbidden_strings(doc: Document) -> list[str]:
    errors: list[str] = []
    text = "\n".join(p.text for p in doc.paragraphs)
    for needle in FORBIDDEN_STRINGS:
        if needle in text:
            errors.append(f"Forbidden string remains in manuscript: {needle}")
        else:
            ok(f"Forbidden string absent: {needle}")
    return errors


def check_figure7(doc: Document) -> list[str]:
    errors: list[str] = []
    hits = [i for i, p in enumerate(doc.paragraphs[:246]) if "Figure 7" in p.text]
    if hits:
        errors.append(f"Figure 7 still cited or labeled in body at paragraphs {hits}")
    else:
        ok("No Figure 7 reference remains in manuscript body")
    return errors


def check_supp_refs(doc: Document) -> list[str]:
    errors: list[str] = []
    body_texts = [p.text for p in doc.paragraphs[:246]]
    for ref in EXPECTED_SUPP_REFS:
        hits = [i for i, text in enumerate(body_texts) if ref in text]
        if not hits:
            errors.append(f"Supplementary item not cited in main body: {ref}")
        else:
            ok(f"{ref} cited in body at paragraphs {hits}")
    return errors


def check_placeholder_table_caption(doc: Document) -> list[str]:
    errors: list[str] = []
    if len(doc.paragraphs) > 254 and doc.paragraphs[254].text.strip():
        errors.append("Placeholder paragraph at 254 is not empty")
    else:
        ok("Table 1 placeholder paragraph is empty")
    return errors


def check_baseline_numbers() -> list[str]:
    errors: list[str] = []
    path = Path("/Volumes/Keywest_JetDrive 1/Codex/Academia/Complaint/tables/bootstrap_results.xlsx")
    ordinary = pd.read_excel(path, sheet_name="ordinary_weeks")
    candidates = pd.read_excel(path, sheet_name="baseline_candidates")
    bootstrap = pd.read_excel(path, sheet_name="bootstrap_summary")

    ordinary_count = int(ordinary["is_ordinary_week"].sum())
    candidate_count = int(len(candidates))

    if ordinary_count != 114:
        errors.append(f"ordinary_weeks count mismatch: expected 114, got {ordinary_count}")
    else:
        ok("ordinary_weeks count = 114")

    if candidate_count != 93:
        errors.append(f"baseline_candidates count mismatch: expected 93, got {candidate_count}")
    else:
        ok("baseline_candidates count = 93")

    expected_significant = {
        ("Flood 2022 - Ordinary", "Anxiety/Worried"): True,
        ("Flood 2022 - Ordinary", "Complaint"): True,
        ("Flood 2022 - Ordinary", "Suspicion/Mistrust"): True,
        ("Flood 2023 - Ordinary", "Anxiety/Worried"): True,
        ("Flood 2023 - Ordinary", "Sad/Disappointed"): True,
        ("Flood 2023 - Ordinary", "Suspicion/Mistrust"): True,
        ("Flood 2023 - Ordinary", "Anger/Rage"): True,
        ("Flood 2023 - Ordinary", "Complaint"): True,
        ("Flood 2023 - Ordinary", "Embarrassment/Distress"): True,
        ("Flood 2023 - Ordinary", "Expectation"): True,
    }
    for (comparison, emotion), expected in expected_significant.items():
        subset = bootstrap.loc[(bootstrap["comparison"] == comparison) & (bootstrap["emotion"] == emotion)]
        if subset.empty:
            errors.append(f"Missing bootstrap result for {comparison} / {emotion}")
            continue
        actual = bool(subset.iloc[0]["significant"])
        if actual != expected:
            errors.append(f"Unexpected significance for {comparison} / {emotion}: got {actual}")
        else:
            ok(f"Bootstrap significance verified for {comparison} / {emotion}")

    return errors


def check_reference_cleanup(doc: Document) -> list[str]:
    errors: list[str] = []
    ref_text = "\n".join(p.text for p in doc.paragraphs[160:246])
    if "ZENODO." in ref_text:
        errors.append("Upper-case ZENODO DOI remains in references")
    else:
        ok("Zenodo DOI case normalized")
    return errors


def main() -> int:
    docx_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_DOCX_PATH
    errors: list[str] = []
    errors.extend(check_files(docx_path))
    if errors:
        for item in errors:
            fail(item)
        return 1

    doc = load_doc(docx_path)
    errors.extend(check_forbidden_strings(doc))
    errors.extend(check_figure7(doc))
    errors.extend(check_supp_refs(doc))
    errors.extend(check_placeholder_table_caption(doc))
    errors.extend(check_baseline_numbers())
    errors.extend(check_reference_cleanup(doc))

    print("")
    if errors:
        for item in errors:
            fail(item)
        fail(f"Preflight check failed with {len(errors)} issue(s)")
        return 1

    ok("Preflight check passed with no blocking issues")
    return 0


if __name__ == "__main__":
    sys.exit(main())
