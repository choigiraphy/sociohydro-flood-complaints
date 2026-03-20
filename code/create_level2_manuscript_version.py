from __future__ import annotations

from pathlib import Path

from docx import Document


SOURCE_PATH = Path("/Volumes/Keywest_JetDrive 1/학교/2026-1/AGU_WRR_Choi/최신_AGU_Manuscript_Choi_수정중.docx")
OUTPUT_PATH = Path("/Volumes/Keywest_JetDrive 1/학교/2026-1/AGU_WRR_Choi/최신_AGU_Manuscript_Choi_수정중_Level2_20260320_v16.docx")


REPLACEMENTS = {
    9: (
        "Civil complaints lag extreme rainfall by 1 to 3 days, indicating later institutional response than rapid search activity."
    ),
    10: (
        "Complaint emotions differed by event: 2022 was anxiety-dominant, whereas 2023 showed broader negative emotional shifts."
    ),
    11: (
        "Rainfall-complaint spatial alignment was weaker in 2023, suggesting response patterns were not explained by exposure alone."
    ),
    13: (
        "Understanding societal response to floods is important for integrated risk assessment, yet upper-tail rainfall events can generate different formal public reactions. We analyzed daily precipitation, internet search activity, and administrative civil complaints during two major South Korean floods in 2022 and 2023 to examine temporal synchronization, emotional contrasts, and spatial rainfall-response coupling. Cross-correlation and block-bootstrap analyses showed a consistent temporal ordering: search activity responded near rainfall peaks, whereas civil complaints lagged by 1 to 3 days. Relative to a non-flood baseline, the 2022 event was characterized mainly by elevated anxiety, while the 2023 event showed broader increases in mistrust, anger, complaint, sadness, and distress. Spatial analyses further showed weaker rainfall-complaint alignment in 2023 than in 2022. National media framing in 2023 was more governance-oriented, but this is treated only as contextual evidence. These results indicate that administrative complaints provide a useful indicator of institutionalized civic response that complements hazard and attention measures."
    ),
    14: "",
    15: "",
    16: "",
    28: (
        "Administrative civil complaints address this gap. Unlike social posts or anonymous expressions, complaints require procedural "
        "effort and direct contact with public agencies. They therefore provide an observable form of institutionalized civic response, "
        "reflecting where hydrological impacts intersect with infrastructure capacity, service disruption, and governance expectations."
    ),
    29: (
        "Risk communication research distinguishes hazard magnitude from outrage, where outrage reflects trust, control, and accountability "
        "rather than physical exposure alone (Sandman, 1993). This distinction motivates our comparative design (Sandman, 1993; "
        "Di Baldassarre et al., 2019). The purpose of the comparison is therefore not to treat the two events as hydrologically identical, "
        "but to examine whether institutionalized civic response behaves similarly across two upper-tail flood episodes of comparable duration. "
        "If civic response were driven primarily by exposure, stronger hydrological forcing would be expected to produce stronger coupling "
        "between rainfall and complaint response. If attribution context also matters, upper-tail events may yield different emotional regimes "
        "and spatial patterns of institutional engagement."
    ),
    30: (
        "We therefore integrate precipitation records, internet search trends, and administrative civil complaints across two major flood "
        "events in South Korea (2022 and 2023) and ask: (1) how are rainfall, attention, and complaints temporally synchronized during "
        "flood windows? (2) do complaint emotions differ across events occurring under overlapping extreme conditions? and (3) how does the "
        "spatial distribution of complaint-based indicators correspond to rainfall exposure across regions? By bringing administrative "
        "complaints into this comparison, the study extends socio-hydrological analysis from exposure and attention signals to formal "
        "institutional response and operationalizes that response using observable complaint timing, emotional composition, and regional distribution."
    ),
    39: (
        "Although the 2023 event exhibited higher RX1day and accumulated precipitation than the 2022 event, both floods occurred under rare, "
        "high-impact precipitation conditions. Our analysis therefore focuses on how societal response regimes diverge across two 14-day upper-tail "
        "flood episodes, rather than attributing differences solely to magnitude. Comparability here refers to shared event class "
        "(rare, high-impact, short-duration flood episodes) and common analytical framing, not to equality of hydrological magnitude or impact "
        "history (Supplementary Table S1)."
    ),
    46: (
        "To capture qualitative dimensions of civic response, emotional content was inferred using a supervised Korean emotion classification "
        "model based on the KOTE taxonomy. Emotion probabilities were assigned at the individual complaint level and aggregated by event window "
        "and region. While individual-level predictions contain uncertainty, aggregation enables robust comparison of event-level emotional "
        "regimes (Jeon et al., 2022; Supplementary Figure S1)."
    ),
    49: (
        "Statistical significance was assessed using a block-bootstrap null model following established practices in time-series analysis of "
        "short event windows (Efron & Tibshirani, 1993; Davison & Hinkley, 1997). We randomly sampled 5,000 non-flood 14-day blocks from the "
        "2021-2023 record. For each resampled block, lagged correlations were recomputed using the same procedure, generating empirical null "
        "distributions at each lag. Empirical p-values were calculated as the proportion of null correlations exceeding the observed flood-period "
        "correlation."
    ),
    50: (
        "This approach accounts for serial dependence and avoids assumptions of independence or normality in short event windows "
        "(Supplementary Figure S2)."
    ),
    57: (
        "Spatial coupling between precipitation intensity and civic response indicators was evaluated using Spearman rank correlations "
        "(n = 17), which provide robustness under small-sample conditions and non-normal distributions. Alternative spatial diagnostics "
        "are reported in Supplementary Figure S4."
    ),
    59: (
        "To quantify departures from ordinary institutional activity, baseline periods were constructed in two steps. "
        "First, we excluded all weeks overlapping the two focal flood events, defined as the 2022 Seoul flood "
        "(August 8-21, 2022) and the 2023 Osong flood (July 10-23, 2023). Among the remaining non-flood weeks, "
        "we defined ordinary weeks as those with weekly complaint counts less than or equal to 40."
    ),
    60: (
        "This threshold was selected empirically from the distribution of non-flood complaint activity to isolate "
        "the low-to-moderate complaint regime representing routine background variation rather than event-driven surges. "
        "In the non-flood record, weekly complaint counts had a mean of 32.2 and a median of 31.0, whereas the subset "
        "satisfying the ordinary-week criterion had a mean of 26.9 and a median of 26.5. We therefore interpreted weeks "
        "at or below this threshold as representing normal civic complaint conditions. The sampling unit for the baseline "
        "was a 14-day block, defined as two consecutive ordinary weeks. Candidate baseline blocks were retained only when "
        "both adjacent weeks satisfied the ordinary-week criterion, and bootstrap resampling with replacement was then "
        "performed from this candidate set. This design preserves duration comparability with the two-week flood windows "
        "while avoiding artificial combinations of temporally disconnected weeks."
    ),
    53: (
        "Differences between events were evaluated using bootstrap resampling. For flood-versus-baseline comparisons, complaints were "
        "resampled with replacement within each flood event window, while the ordinary baseline was resampled at the level of eligible "
        "14-day blocks, treating each baseline block as an independent observational unit. Event-to-event comparisons were evaluated by "
        "resampling complaints within each flood window. Confidence intervals were derived from the resulting bootstrap distributions, "
        "allowing nonparametric inference while preserving the ordinary-period block structure."
    ),
    71: (
        "Figure 5 summarizes baseline-referenced differences in the emotional composition of civil complaints across "
        "the two flood events."
    ),
    72: (
        "Relative to the revised baseline pool, the 2022 flood showed a significant increase in Anxiety/Worried "
        "(mean difference = 0.145, 95% CI [0.106, 0.184]) and significant decreases in Complaint "
        "(mean difference = -0.059, 95% CI [-0.105, -0.012]) and Suspicion/Mistrust "
        "(mean difference = -0.060, 95% CI [-0.091, -0.029])."
    ),
    73: (
        "In contrast, the 2023 flood exhibited significant increases in Anxiety/Worried "
        "(0.110, [0.070, 0.149]), Sad/Disappointed (0.082, [0.046, 0.120]), Suspicion/Mistrust "
        "(0.067, [0.027, 0.107]), Anger/Rage (0.063, [0.012, 0.113]), Complaint (0.064, [0.014, 0.114]), "
        "and Embarrassment/Distress (0.029, [0.004, 0.052]), together with a decrease in Expectation "
        "(-0.036, [-0.064, -0.009])."
    ),
    74: (
        "These recalculated results preserve the core event contrast under a block-based ordinary baseline: the 2022 event is characterized "
        "primarily by elevated anxiety relative to ordinary conditions, whereas the 2023 event shows a broader negative emotional profile "
        "that includes mistrust, anger, complaint, sadness, and distress. These bootstrap results provide the direct evidence for the "
        "event-level emotional contrast discussed below. Between-event bootstrap comparisons likewise indicate that "
        "Sad/Disappointed, Anger/Rage, Complaint, and Suspicion/Mistrust were higher in 2023 than in 2022."
    ),
    76: (
        "Figure 3 evaluates whether regions with higher rainfall exposure also showed higher complaint intensity during each event, whereas "
        "Figure 6 localizes baseline-referenced shifts in selected negative emotions."
    ),
    77: (
        "At the complaint-intensity level, the 2022 event showed closer spatial correspondence between precipitation and complaints: regions "
        "with higher rainfall tended to coincide more closely with elevated complaint intensity. In contrast, the 2023 event showed weaker "
        "alignment, indicating that complaint intensity was not distributed in proportion to rainfall exposure alone."
    ),
    80: "3.4 Departure from Ordinary Periods",
    81: (
        "Figure 4 compares complaint activity during flood windows with ordinary (non-flood) conditions, whereas "
        "Figures 5 and 6 evaluate how emotional composition and regional emotion shifts depart from the revised baseline definition."
    ),
    82: (
        "Under the final specification, the baseline consisted of 114 ordinary weeks and 93 eligible consecutive two-week "
        "baseline blocks. Non-flood weekly complaint counts averaged 32.2 complaints per week, whereas ordinary weeks averaged "
        "26.9 complaints per week. The resulting candidate 14-day baseline blocks had a mean total of 52.1 complaints "
        "(interquartile range: 43-63), indicating that the baseline captures low-to-moderate non-disaster complaint activity "
        "rather than event-driven peaks."
    ),
    83: (
        "Relative to this baseline, emotional departures remain event-specific. Anxiety-related expressions rose above baseline "
        "during both flood events, but 2023 additionally showed statistically significant increases in complaint, mistrust, anger, "
        "sadness, and embarrassment/distress. These contrasts support the interpretation that the two events differed not simply "
        "in magnitude, but in the emotional structure of institutionalized civic response."
    ),
    85: (
        "To assess robustness, precipitation was aggregated using alternative rolling accumulation windows, and lagged correlation "
        "analyses were repeated as a sensitivity analysis."
    ),
    62: (
        "Several robustness checks were performed. Temporal analyses were repeated using alternative precipitation aggregation windows to "
        "confirm that the ordering of response signals was not sensitive to specification (Supplementary Figure S3). Spatial analyses were "
        "tested using alternative normalization choices, and ordinary-week threshold sensitivity is summarized in Supplementary Table S5."
    ),
    67: (
        "In both events, search activity responded rapidly to precipitation peaks, exhibiting near-zero lag. In contrast, civil complaints "
        "consistently lagged precipitation by several days. Lagged Pearson correlation analysis, following standard hydrological time-series "
        "methods (Box et al., 2015), identified the strongest precipitation-complaint association at +1 day in 2022 (r = 0.728, n = 13 "
        "effective pairs) and at +3 days in 2023 (r = 0.796, n = 11 effective pairs)."
    ),
    69: (
        "These results indicate a robust positive-lag relationship between precipitation and institutionalized civic response, with event-dependent differences in lag duration but consistent ordering of response signals."
    ),
    78: (
        "This difference is supported by the spatial statistics. In 2022, negative-emotion indicators showed only weak to moderate spatial "
        "autocorrelation (I = 0.22-0.38, p = 0.08-0.14, depending on emotion category). In 2023, mistrust-related complaints exhibited stronger "
        "and statistically significant clustering (I = 0.45, p = 0.002), indicating that negative institutional response became more spatially concentrated."
    ),
    79: (
        "Spearman correlations further distinguish these patterns. In 2023, the correlation between precipitation intensity and mistrust-related "
        "complaint proportions was weak and statistically insignificant (ρ = 0.12, p = 0.61), indicating that the strongest mistrust signal did "
        "not occur simply where rainfall was highest. Taken together, Figure 3 and the spatial statistics provide the direct evidence for weaker exposure-response coupling "
        "in 2023 than in 2022."
    ),
    88: "3.6 Regional Heterogeneity in Response Magnitude (Figure 6)",
    89: (
        "Figure 6 illustrates regional variation in baseline-referenced emotion shifts for three focal emotions "
        "(Anger/Rage, Sad/Disappointed, and Embarrassment/Distress)."
    ),
    90: (
        "Regional emotion shifts were mapped as flood-period dominant-emotion proportions minus the corresponding regional baseline proportions "
        "derived from the bootstrap-sampled baseline pool. These maps are intended to identify where negative emotions intensified relative to "
        "ordinary regional conditions, rather than to measure rainfall-complaint coupling directly. To avoid unstable estimates from sparse counts, "
        "maps were restricted to regions with at least five complaints in both the flood-period sample and the regional baseline sample."
    ),
    91: (
        "Under this criterion, eight provinces were retained for each mapped emotion in 2022 and thirteen in 2023. In 2022, anger shifts were "
        "positive mainly in Seoul, Daejeon, and Gyeongbuk, while sadness increases were strongest in Chungbuk, Chungnam, and Gyeonggi; "
        "embarrassment/distress remained near zero or slightly negative across most mapped provinces. In 2023, anger shifts became more spatially "
        "extensive, with the largest positive values in Chungbuk and Gwangju, whereas sadness showed mixed regional signs and "
        "embarrassment/distress remained modest except for localized increases in Gyeongnam and Gwangju. Figure 6 therefore complements Figure 3 "
        "by showing where negative emotional intensification was localized after baseline adjustment."
    ),
    95: (
        "At the national level, the proportion of institutional-failure framing increased from 5.35% in 2022 to 9.62% in 2023. Although media "
        "data do not provide region-specific attribution signals, this event-level shift is directionally consistent with the stronger clustering "
        "of negative complaint emotions and the weaker rainfall-complaint coupling observed in 2023. These media results are used as contextual, "
        "rather than inferential, evidence. Keyword rules for the framing classification are "
        "reported in Supplementary Table S2."
    ),
    170: (
        "Bell, S. W. (2020). samwbell/saturn_counts (Version 1.1.0) [Software]. Zenodo. https://doi.org/10.5281/zenodo.3766959"
    ),
    172: (
        "Box, G. E., Jenkins, G. M., Reinsel, G. C., & Ljung, G. M. (2015). Time Series Analysis: Forecasting and Control. John Wiley & Sons."
    ),
    186: (
        "Efron, B., & Tibshirani, R. J. (1993). An Introduction to the Bootstrap. Chapman & Hall."
    ),
    205: (
        "Kasperson, R. E., Renn, O., Slovic, P., et al. (1988). The social amplification of risk. Risk Analysis, 8(2), 177-187."
    ),
    209: (
        "Kim, J., Lee, S., & Park, H. (2023). Hydrological analysis of the August 2022 extreme rainfall in Seoul. Journal of Korea Water Resources Association, 56(3), 187-201. [In Korean, English abstract]"
    ),
    215: (
        "Lee, J. (2021). A structured self-attentive sentence embedding. arXiv:1703.03130."
    ),
    229: (
        "Park, J.-S., & Lee, S.-M. (2020). Big data analysis of civil affairs. Informatization Policy, 27(2), 66-83. [In Korean]"
    ),
    235: (
        "Sandman, P. M. (1993). Responding to Community Outrage. American Industrial Hygiene Association."
    ),
    243: (
        "World Bank. (2022). Grievance Redress Service: Annual Report FY22. World Bank."
    ),
    250: (
        "Figure 3. Spatial distribution of precipitation intensity and complaint indicators by first-order administrative region. (a-b) "
        "Population-normalized complaint intensity during the 2022 and 2023 flood windows. (c-d) Corresponding regional precipitation intensity. "
        "This figure is used to assess event-specific spatial alignment between rainfall exposure and complaint intensity: alignment is closer in "
        "2022 and weaker in 2023."
    ),
    252: (
        "Figure 5. Bootstrap-based emotion shifts during the 2022 and 2023 flood periods relative to the non-flood baseline. "
        "The baseline was defined as bootstrap-sampled 14-day blocks composed of two consecutive ordinary weeks, where ordinary "
        "weeks were non-flood weeks with weekly complaint counts less than or equal to 40. Points indicate mean differences in "
        "emotion proportion relative to the baseline, and horizontal error bars indicate 95% bootstrap confidence intervals obtained by "
        "resampling event complaints and ordinary-period baseline blocks. "
        "Only emotions with confidence intervals excluding zero in at least one flood period are shown."
    ),
    253: (
        "Figure 6. Regional shifts in dominant-emotion proportions during the flood periods relative to the bootstrap-derived baseline. "
        "For each province and emotion, the mapped value is the flood-period dominant-emotion proportion minus the corresponding "
        "proportion in the regional baseline complaint pool. The baseline pool was constructed from bootstrap-sampled 14-day blocks "
        "of two consecutive ordinary weeks. This figure localizes baseline-referenced emotional intensification and is not interpreted "
        "as a direct rainfall-coupling measure. Hatched regions indicate insufficient data (fewer than five complaints in either the "
        "flood-period sample or the regional baseline sample)."
    ),
    254: "",
    289: (
        "Supplementary Figure S2 illustrates the baseline and block-bootstrap procedure used for temporal synchronization and baseline "
        "comparisons. Ordinary weeks were defined as non-flood weeks with weekly complaint counts less than or equal to 40, and baseline "
        "units were constructed as 14-day blocks composed of two consecutive ordinary weeks. Bootstrap resampling was then performed from "
        "this candidate set of observed non-flood blocks. This design preserves duration comparability with the flood windows while "
        "avoiding artificial combinations of temporally disconnected weeks."
    ),
    295: (
        "Supplementary Figure S4 presents additional spatial diagnostics, including alternative spatial weights matrices and "
        "robustness checks for Moran’s I statistics. Results are consistent with the main text, confirming event-dependent differences "
        "in spatial coupling between rainfall intensity and complaint emotions. Supplementary Table S5 reports threshold sensitivity "
        "checks for the ordinary-week definition using weekly complaint cutoffs of 30, 40, and 50. Across these alternatives, the "
        "number of significant flood-versus-baseline emotion differences remained stable for the 2022 event and varied only marginally "
        "for the 2023 event, indicating that the substantive conclusions are not driven by the exact threshold choice."
    ),
    275: "",
    97: "4.1 Temporal Structure of Institutionalized Societal Response",
    98: (
        "The results indicate a consistent temporal ordering of societal response signals during both extreme rainfall events. "
        "In each case, precipitation peaks were followed quickly by increases in information-seeking behavior, whereas civil complaints "
        "emerged after a delay of one to several days."
    ),
    99: (
        "This separation suggests that the available data streams capture different stages of societal response. Search activity appears "
        "to reflect rapid awareness and attention, whereas civil complaints represent a later stage of formal engagement with public "
        "institutions. Because the lag structure was preserved across events and alternative aggregation choices, this temporal distinction is directly supported by the observed time-series structure and is unlikely to be "
        "explained solely by smoothing or indicator selection."
    ),
    100: (
        "Accordingly, the findings suggest that institutionalized civic response should not be treated as interchangeable with real-time "
        "hazard intensity or with rapid attention metrics. Instead, complaints appear to capture a later stage in which hydrological "
        "impacts are processed into formal demands for institutional response."
    ),
    102: (
        "Despite both events occurring under extreme rainfall conditions, the emotional composition of civil complaints differed "
        "substantially between 2022 and 2023. Under the revised baseline definition, the 2022 event was characterized mainly by elevated "
        "anxiety, whereas the 2023 event showed simultaneous increases in mistrust, anger, complaint, sadness, and distress."
    ),
    103: (
        "This contrast does not align with a simple monotonic expectation based on hydrological magnitude alone. The 2023 event was not "
        "associated with stronger rainfall-response coupling across all dimensions; instead, it combined stronger negative emotional shifts "
        "with weaker spatial correspondence between rainfall and complaint response."
    ),
    104: (
        "A cautious interpretation is therefore that upper-tail flood events of comparable analytical duration can be associated with different "
        "emotional response regimes. The present results are consistent with socio-hydrological accounts in which civic response depends not only "
        "on exposure, but also on how impacts are interpreted within institutional and communicative contexts. Here, the direct evidence comes from "
        "the bootstrap emotion contrasts and the spatial coupling analyses; the broader socio-hydrological reading is an interpretation of those results rather than a separate empirical test."
    ),
    106: (
        "The spatial analyses support an event-dependent relationship between hydrological exposure and civic response. Three observations are "
        "most relevant: Figure 3 shows closer rainfall-complaint intensity alignment in 2022 than in 2023; the spatial statistics show stronger "
        "clustering of mistrust-related complaints in 2023; and Figure 6 indicates that negative-emotion intensification in 2023 was localized "
        "across a broader set of regions after baseline adjustment."
    ),
    107: (
        "Taken together, these results support a description of relative spatial decoupling in 2023, rather than a claim that civic response "
        "became fully independent of exposure. Rainfall remained an important initiating condition, but it did not by itself explain either the "
        "regional distribution of complaint intensity or the regions showing the strongest negative emotional shifts."
    ),
    108: (
        "Figure 3 therefore identifies where exposure-response coupling weakens at the complaint-intensity level, whereas Figure 6 shows that, "
        "once compared against ordinary regional baselines, the emotional response was not spatially uniform. Read together, the two figures "
        "support a bounded claim: the translation from exposure into institutionalized response varied across events and regions."
    ),
    110: (
        "Independent media analysis provides contextual evidence that is directionally consistent with this interpretation. National news coverage "
        "during the 2023 event contained a higher share of governance-oriented framing than during the 2022 event, placing greater emphasis on "
        "institutional responsibility and response failure."
    ),
    111: (
        "Because the media data are available only at the national level, they cannot explain the regional complaint pattern directly. "
        "However, the shift in framing is directionally consistent with the event-level increase in mistrust-related complaint emotions and "
        "the weaker rainfall-complaint spatial coupling observed in 2023."
    ),
    112: (
        "Media framing is therefore used here as contextual evidence rather than as a causal explanatory variable. The combined pattern is "
        "consistent with, but does not by itself establish, an attribution-mediated pathway linking hydrological impacts to formal civic response."
    ),
    114: (
        "These findings indicate that administrative civil complaints provide information that is not captured fully by hazard metrics or "
        "rapid attention indicators alone. Precipitation and search activity are useful for tracking exposure and awareness, whereas "
        "complaints provide a distinct view of when citizens escalate from concern to formal institutional engagement. The contribution of this "
        "study is therefore not only descriptive; it is to show that formal complaint records can be operationalized as an analyzable socio-hydrological signal. "
        "That contribution is supported directly by the temporal, emotional, and spatial results reported above."
    ),
    115: (
        "From a socio-hydrological perspective, this expands the observable range of human-water interactions by linking one physical signal "
        "(precipitation), one rapid attention signal (search behavior), and one formal institutional-response signal (administrative complaints). "
        "Temporal lags, emotional composition, and regional clustering together provide complementary evidence on how hydrological events are "
        "translated into institutional response."
    ),
    116: (
        "Operationally, such multi-signal approaches may help distinguish cases in which a flood functions primarily as a physical hazard "
        "from those in which it also becomes a governance and communication problem. The present two-event design does not establish a general "
        "law, but it does show that this distinction can be measured empirically and compared across events."
    ),
    118: (
        "Several limitations define the scope of inference. First, the analysis is based on two flood events within a single national context. "
        "Accordingly, the study should be read as an event-comparative test of whether institutionalized civic response can vary across upper-tail "
        "flood episodes, not as an estimate of population-level regularities across all flood types or governance settings."
    ),
    119: (
        "Second, the hydrological comparison is framed at the level of event class and analytical duration rather than exact equality of "
        "meteorological forcing, exposure, or damage history. This limits any claim that the observed response differences can be attributed to "
        "a single causal factor, but it does not undermine the narrower result that response timing, emotional composition, and spatial "
        "organization differed across two rare, high-impact 14-day flood episodes."
    ),
    120: (
        "Third, emotion classification relies on supervised text models and therefore carries uncertainty at the individual complaint level. "
        "The current design reduces this risk by aggregating predictions and by focusing inference on event-level and regional contrasts, but "
        "the results should still be interpreted as estimates of comparative emotional structure rather than precise measurements of latent affect."
    ),
    121: (
        "Fourth, the spatial analysis is conducted at the first-order administrative-region level, which necessarily smooths subregional "
        "variation in inundation, infrastructure failure, and complaint generation. The spatial results therefore support claims about "
        "regional alignment and localization, not parcel-scale matching between hazard exposure and individual response."
    ),
    122: (
        "Fifth, the media-framing analysis is national in scale and is used only as contextual evidence. It cannot identify regional attribution "
        "mechanisms or establish that media narratives caused the observed complaint patterns. More localized media archives, institutional "
        "performance indicators, or interview-based evidence would be needed to test attribution pathways directly."
    ),
    123: (
        "Within these scope conditions, the manuscript's main claims remain narrower and more defensible: formal complaints lag rapid attention "
        "signals, emotional complaint structure differs across the two events, and rainfall-complaint coupling was weaker in 2023 than in 2022."
    ),
    124: (
        "Although the empirical setting is South Korea, the broader implication is methodological rather than universal. Many governance systems "
        "maintain formal grievance or complaint channels, so the same analytic logic can be evaluated elsewhere, but transferability should be "
        "established through additional cases rather than assumed."
    ),
    127: (
        "This study examined how extreme rainfall events translate into societal responses by integrating precipitation records, internet search "
        "behavior, and administrative civil complaints across two major flood events in South Korea. By jointly analyzing temporal dynamics, "
        "emotional composition, and spatial patterns, the study moves beyond hazard-centric assessments to characterize when and where hydrological "
        "forcing is associated with institutionalized civic response."
    ),
    128: (
        "Three main conclusions emerge. First, societal responses follow a consistent temporal structure under extreme rainfall. Information-seeking "
        "behavior responds almost immediately to precipitation peaks, whereas civil complaints emerge after a systematic delay of several days. "
        "This lagged structure is robust across events and analytical specifications, indicating that complaints capture a later, more deliberative "
        "stage of engagement distinct from rapid public attention."
    ),
    129: (
        "Second, the two upper-tail flood episodes were associated with different emotional response regimes. Complaints during the 2022 event were "
        "dominated mainly by anxiety-related expressions, whereas the 2023 event showed broader increases in mistrust-, anger-, complaint-, sadness-, "
        "and distress-related indicators. This contrast suggests that complaint emotions were not ordered solely by hydrological severity."
    ),
    130: (
        "Third, the spatial coupling between rainfall intensity and civic response was event-dependent. In 2022, complaint activity more closely "
        "aligned with high-rainfall regions, while in 2023, mistrust-related complaint clusters were less tightly aligned with the highest "
        "precipitation zones. This pattern indicates that flood-related societal impacts cannot be inferred from hydrological metrics alone."
    ),
    131: (
        "Together, these findings identify administrative civil complaints as a useful complementary signal in socio-hydrological analysis. Complaints "
        "capture institutionalized civic response that is not fully reflected in digital attention indicators, providing a measurable way to compare "
        "how hydrological extremes are translated into formal public response."
    ),
    132: (
        "The interpretation should remain bounded by the study design. The analysis is based on two events within a single national context, and the "
        "emotion and media analyses are subject to the scale and uncertainty constraints described above."
    ),
    133: (
        "Future work could extend this framework to additional flood events, incorporate hydrodynamic modeling outputs, and examine feedbacks between "
        "repeated exposure and evolving institutional response. Incorporating higher-resolution spatial data and infrastructure characteristics would "
        "clarify when event-dependent coupling and decoupling are most likely to emerge."
    ),
    19: (
        "Floods are usually evaluated by rainfall intensity and physical damage, but formal public responses do not always follow those measures. We compared two major South Korean flood events using rainfall data, internet searches, and civil complaints submitted to government agencies. People searched for flood-related information almost immediately after heavy rain, but complaints were filed 1 to 3 days later, suggesting a later and more deliberate form of response. The 2022 event was marked mainly by anxiety, whereas the 2023 event showed broader negative emotions such as mistrust and anger. Complaints were also less tightly aligned with the regions of highest rainfall in 2023. Administrative complaint records can therefore complement hazard and online attention data when evaluating how flood events become broader social and governance problems."
    ),
    20: "",
    21: "",
    44: (
        "Administrative civil complaints were obtained from South Korea’s national civil-complaint management system. Each record includes submission date, administrative region, and complaint text. The study analyzed secondary administrative records supplied under the MOU with ACRC and did not involve direct interaction with human participants. Flood-related complaints were identified through keyword filtering and manual validation."
    ),
    136: (
        "This research was supported by the Institute of Social Data Science (ISDS) at POSTECH and conducted under a Memorandum of Understanding with South Korea's Anti-Corruption and Civil Rights Commission (ACRC). We thank the ACRC for data access and the KMA for meteorological data."
    ),
    140: (
        "The civil complaint data were obtained under a Memorandum of Understanding (MOU) with the Anti-Corruption and Civil Rights Commission (ACRC) of South Korea. Due to privacy regulations under the Personal Information Protection Act, raw complaint texts are not publicly available. Access to the restricted source records is controlled by ACRC."
    ),
    142: (
        "News keyword data were retrieved from BigKinds (https://www.bigkinds.or.kr), operated by the Korea Press Foundation. Publicly shareable, de-identified aggregated datasets supporting the analyses and figures are archived at Zenodo: https://doi.org/10.5281/zenodo.19133036."
    ),
    145: (
        "All code required to reproduce the analytical workflow, figures, and summary tables is archived at Zenodo: https://doi.org/10.5281/zenodo.19133036 and mirrored on GitHub at https://github.com/choigiraphy/sociohydro-flood-complaints. Where raw data cannot be shared due to legal restrictions, we provide aggregated inputs and fully documented scripts that reproduce the analytical logic and parameterization used in the manuscript."
    ),
    146: (
        "The reproducibility package includes:"
    ),
    152: (
        "Recommended environment: Python 3.9, PyTorch 1.12, Transformers 4.21, and Geopandas 0.11."
    ),
    155: (
        "The authors declare there are no conflicts of interest for this manuscript."
    ),
    159: (
        "Choi, G., Kam, J., Bae, Y., & Shin, H. (2026). Sociohydro flood complaints reproducibility package (Version 1.0.0) [Software]. Zenodo. https://doi.org/10.5281/zenodo.19133036"
    ),
    164: "",
    166: "",
    168: "",
    170: "",
    174: "",
    187: "",
    189: "",
    197: "",
    209: "",
    215: "",
    217: "",
    227: "",
    229: "",
    245: "",
}


def main() -> None:
    doc = Document(SOURCE_PATH)
    for idx, new_text in REPLACEMENTS.items():
        if idx >= len(doc.paragraphs):
            raise IndexError(f"Paragraph index {idx} is out of range for {SOURCE_PATH.name}")
        doc.paragraphs[idx].text = new_text
    doc.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
